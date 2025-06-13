"""
Document Parser with Enhanced Voice Support
Multiple voice input options for Windows compatibility
"""

import streamlit as st
import io
import tempfile
import os
import base64

# Voice recognition imports with fallbacks
VOICE_METHOD = None
VOICE_AVAILABLE = False

# Try different audio recording packages
try:
    from audio_recorder_streamlit import audio_recorder
    VOICE_METHOD = "audio_recorder_streamlit"
    VOICE_AVAILABLE = True
except ImportError:
    try:
        from streamlit_webrtc import webrtc_streamer, AudioProcessorBase, WebRtcMode
        import av
        VOICE_METHOD = "webrtc"
        VOICE_AVAILABLE = True
    except ImportError:
        try:
            import st_audiorec as sar
            VOICE_METHOD = "st_audiorec"
            VOICE_AVAILABLE = True
        except ImportError:
            VOICE_METHOD = None
            VOICE_AVAILABLE = False

# Speech recognition
try:
    import speech_recognition as sr
    SPEECH_RECOGNITION_AVAILABLE = True
except ImportError:
    SPEECH_RECOGNITION_AVAILABLE = False

# Document parsing imports
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def check_voice_capabilities():
    """Check available voice input methods"""
    capabilities = {
        'voice_method': VOICE_METHOD,
        'voice_available': VOICE_AVAILABLE,
        'speech_recognition': SPEECH_RECOGNITION_AVAILABLE,
        'audio_recorder_streamlit': VOICE_METHOD == "audio_recorder_streamlit",
        'webrtc': VOICE_METHOD == "webrtc",
        'st_audiorec': VOICE_METHOD == "st_audiorec"
    }
    return capabilities

def install_voice_instructions():
    """Display installation instructions for voice features"""
    st.info("🎤 **Voice features not available**. Install one of these packages:")
    
    with st.expander("📦 Voice Installation Options", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Option 1: audio-recorder-streamlit**")
            st.code("pip install streamlit-audio-recorder\npip install SpeechRecognition")
            
            st.write("**Option 2: streamlit-webrtc**")
            st.code("pip install streamlit-webrtc\npip install SpeechRecognition")
        
        with col2:
            st.write("**Option 3: st-audiorec**")
            st.code("pip install st-audiorec\npip install SpeechRecognition")
            
            st.write("**For Windows PyAudio issues:**")
            st.code("pip install pipwin\npipwin install pyaudio")

def parse_pdf(uploaded_file):
    """Parse PDF file and extract text content."""
    if not PDF_AVAILABLE:
        return "Error: PyPDF2 not installed. Please install: pip install PyPDF2"
    
    try:
        if hasattr(uploaded_file, 'read'):
            content = uploaded_file.read()
            uploaded_file.seek(0)
            pdf_file = io.BytesIO(content)
        else:
            pdf_file = uploaded_file
            
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        result = text.strip()
        return result if result else "No text could be extracted from this PDF."
        
    except Exception as e:
        return f"Error parsing PDF: {e}"

def parse_docx(uploaded_file):
    """Parse DOCX file and extract text content."""
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Please install: pip install python-docx"
    
    try:
        if hasattr(uploaded_file, 'read'):
            content = uploaded_file.read()
            uploaded_file.seek(0)
            docx_file = io.BytesIO(content)
        else:
            docx_file = uploaded_file
            
        doc = Document(docx_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        result = text.strip()
        return result if result else "No text could be extracted from this document."
        
    except Exception as e:
        return f"Error parsing DOCX: {e}"

def parse_txt(uploaded_file):
    """Parse TXT file and extract text content."""
    try:
        if hasattr(uploaded_file, 'read'):
            content = uploaded_file.read()
            uploaded_file.seek(0)
        else:
            content = uploaded_file
            
        if isinstance(content, bytes):
            encodings = ['utf-8', 'utf-8-sig', 'ascii', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            text = content.decode('utf-8', errors='ignore')
            return text.strip()
        else:
            return str(content).strip()
            
    except Exception as e:
        return f"Error parsing TXT: {e}"

def transcribe_audio(audio_bytes, method="google"):
    """
    Transcribe audio to text using speech recognition
    
    Args:
        audio_bytes: Raw audio bytes
        method: Recognition method ("google", "sphinx", "wit")
    
    Returns:
        Transcribed text
    """
    if not SPEECH_RECOGNITION_AVAILABLE:
        return "Error: SpeechRecognition not installed. Please install: pip install SpeechRecognition"
    
    try:
        # Save audio to temporary file
        temp_dir = tempfile.gettempdir()
        temp_filename = f"temp_audio_{os.getpid()}.wav"
        temp_path = os.path.join(temp_dir, temp_filename)
        
        with open(temp_path, "wb") as temp_file:
            temp_file.write(audio_bytes)
        
        # Initialize recognizer
        recognizer = sr.Recognizer()
        
        try:
            # Load and process audio
            with sr.AudioFile(temp_path) as source:
                recognizer.adjust_for_ambient_noise(source, duration=0.5)
                audio = recognizer.record(source)
            
            # Try different recognition methods
            recognition_methods = []
            
            if method == "google" or method == "auto":
                recognition_methods.append(
                    ("Google Web Speech", lambda: recognizer.recognize_google(audio, language='en-US'))
                )
            
            if method == "sphinx" or method == "auto":
                try:
                    recognition_methods.append(
                        ("CMU Sphinx (offline)", lambda: recognizer.recognize_sphinx(audio))
                    )
                except sr.RequestError:
                    pass  # Sphinx not available
            
            # If no specific method or auto, add fallback
            if not recognition_methods or method == "auto":
                recognition_methods.append(
                    ("Google (fallback)", lambda: recognizer.recognize_google(audio))
                )
            
            # Try each method
            for method_name, recognition_func in recognition_methods:
                try:
                    text = recognition_func()
                    if text and text.strip():
                        return text.strip()
                except sr.UnknownValueError:
                    continue
                except sr.RequestError as e:
                    continue
                except Exception as e:
                    continue
            
            return "Could not understand the audio. Please try speaking more clearly or check your microphone."
            
        finally:
            # Clean up temporary file
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except Exception:
                pass
                
    except Exception as e:
        return f"Transcription error: {str(e)}"

def voice_input_interface(key_suffix="", height=None):
    """
    Create voice input interface with multiple options
    
    Args:
        key_suffix: Unique suffix for component keys
        height: Height for text areas
    
    Returns:
        Tuple of (audio_bytes, transcribed_text)
    """
    if not VOICE_AVAILABLE:
        install_voice_instructions()
        return None, ""
    
    audio_bytes = None
    transcribed_text = ""
    
    st.write("🎤 **Voice Input Options:**")
    
    if VOICE_METHOD == "audio_recorder_streamlit":
        st.write("*Using streamlit-audio-recorder*")
        
        # Audio recorder widget
        audio_bytes = audio_recorder(
            text="🎤 Click to Record",
            recording_color="#e87070",
            neutral_color="#6aa36f",
            icon_name="microphone",
            icon_size="2x",
            key=f"audio_recorder_{key_suffix}"
        )
        
        if audio_bytes:
            st.audio(audio_bytes, format="audio/wav")
            
            # Transcription options
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📝 Transcribe (Google)", key=f"transcribe_google_{key_suffix}"):
                    with st.spinner("🔄 Transcribing with Google..."):
                        transcribed_text = transcribe_audio(audio_bytes, "google")
                        if transcribed_text and not transcribed_text.startswith("Error"):
                            st.session_state[f"transcription_{key_suffix}"] = transcribed_text
                            st.success("✅ Transcription successful!")
                        else:
                            st.error(f"❌ {transcribed_text}")
            
            with col2:
                if st.button("🔄 Transcribe (Offline)", key=f"transcribe_offline_{key_suffix}"):
                    with st.spinner("🔄 Transcribing offline..."):
                        transcribed_text = transcribe_audio(audio_bytes, "sphinx")
                        if transcribed_text and not transcribed_text.startswith("Error"):
                            st.session_state[f"transcription_{key_suffix}"] = transcribed_text
                            st.success("✅ Transcription successful!")
                        else:
                            st.error(f"❌ {transcribed_text}")
            
            with col3:
                if st.button("🔄 Record Again", key=f"record_again_{key_suffix}"):
                    if f"transcription_{key_suffix}" in st.session_state:
                        del st.session_state[f"transcription_{key_suffix}"]
                    st.rerun()
    
    elif VOICE_METHOD == "webrtc":
        st.write("*Using streamlit-webrtc*")
        
        # WebRTC audio recorder
        webrtc_ctx = webrtc_streamer(
            key=f"webrtc_{key_suffix}",
            mode=WebRtcMode.SENDONLY,
            audio_receiver_size=1024,
            media_stream_constraints={"video": False, "audio": True},
        )
        
        if webrtc_ctx.audio_receiver:
            st.info("🎤 Recording... Click 'STOP' when finished")
            
            if st.button("📝 Process Recording", key=f"process_webrtc_{key_suffix}"):
                # Process WebRTC audio (implementation depends on your needs)
                st.info("WebRTC audio processing - implementation needed")
    
    elif VOICE_METHOD == "st_audiorec":
        st.write("*Using st-audiorec*")
        
        # st-audiorec widget
        wav_audio_data = sar.st_audiorec()
        
        if wav_audio_data is not None:
            st.audio(wav_audio_data, format='audio/wav')
            
            if st.button("📝 Transcribe Audio", key=f"transcribe_audiorec_{key_suffix}"):
                with st.spinner("🔄 Transcribing..."):
                    transcribed_text = transcribe_audio(wav_audio_data, "auto")
                    if transcribed_text and not transcribed_text.startswith("Error"):
                        st.session_state[f"transcription_{key_suffix}"] = transcribed_text
                        st.success("✅ Transcription successful!")
                    else:
                        st.error(f"❌ {transcribed_text}")
    
    # Show transcription for editing
    if f"transcription_{key_suffix}" in st.session_state:
        st.write("**📝 Edit Transcription:**")
        edited_text = st.text_area(
            "Review and edit the transcribed text:",
            value=st.session_state[f"transcription_{key_suffix}"],
            height=height or 150,
            key=f"edit_transcription_{key_suffix}",
            help="You can edit the transcribed text to fix any errors"
        )
        
        if edited_text != st.session_state[f"transcription_{key_suffix}"]:
            transcribed_text = edited_text
            st.info("📝 Using edited transcription")
        else:
            transcribed_text = st.session_state[f"transcription_{key_suffix}"]
    
    return audio_bytes, transcribed_text

def get_document_input(label: str, doc_type: str, key_suffix: str = ""):
    """
    Enhanced document input with voice support
    
    Args:
        label: Display label for the input
        doc_type: Type of document for context
        key_suffix: Unique suffix for Streamlit keys
        
    Returns:
        Extracted or input content
    """
    upload_key = f"upload_{doc_type}_{key_suffix}"
    text_key = f"text_{doc_type}_{key_suffix}"
    voice_key = f"voice_{doc_type}_{key_suffix}"
    
    # Create tabs for different input methods
    if VOICE_AVAILABLE:
        tab1, tab2, tab3 = st.tabs(["📁 Upload File", "🎤 Voice Input", "💬 Text Input"])
    else:
        tab1, tab3 = st.tabs(["📁 Upload File", "💬 Text Input"])
        tab2 = None
    
    content = ""
    
    # File upload tab
    with tab1:
        st.write(f"**Upload {label}**")
        
        file_types = ['pdf', 'docx', 'txt']
        st.info(f"📄 Supported formats: {', '.join(file_types).upper()}")
        
        uploaded_file = st.file_uploader(
            f"Choose {doc_type} file",
            type=file_types,
            key=upload_key,
            help=f"Upload your {doc_type} document"
        )
        
        if uploaded_file is not None:
            file_size = len(uploaded_file.getvalue())
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("File Size", f"{file_size / 1024:.1f} KB")
            with col2:
                st.metric("File Type", uploaded_file.type)
            
            with st.spinner(f"Processing {uploaded_file.name}..."):
                try:
                    if uploaded_file.type == "application/pdf":
                        content = parse_pdf(uploaded_file)
                    elif uploaded_file.type in [
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "application/msword"
                    ]:
                        content = parse_docx(uploaded_file)
                    elif uploaded_file.type in ["text/plain", "text/txt"]:
                        content = parse_txt(uploaded_file)
                    else:
                        content = parse_txt(uploaded_file)
                    
                    if content and not content.startswith("Error"):
                        st.success(f"✅ Successfully processed {uploaded_file.name}")
                        
                        with st.expander("📋 Content Preview"):
                            preview = content[:300] + "..." if len(content) > 300 else content
                            st.text_area("Preview", preview, height=100, disabled=True)
                            st.write(f"**Length:** {len(content)} characters, {len(content.split())} words")
                    else:
                        st.error(f"❌ {content}")
                        
                except Exception as e:
                    st.error(f"❌ Error processing file: {str(e)}")
    
    # Voice input tab
    if tab2:
        with tab2:
            st.write(f"**Record {label}**")
            st.info("🎤 Record yourself reading or describing the content")
            
            _, transcribed_text = voice_input_interface(voice_key, height=200)
            
            if transcribed_text:
                content = transcribed_text
                st.success("✅ Voice input received!")
    
    # Text input tab
    with tab3:
        st.write(f"**Type {label}**")
        
        placeholders = {
            "job_posting": """Paste the job description here...

Example:
Job Title: Software Engineer
Company: Tech Corp

Responsibilities:
- Develop web applications
- Collaborate with team
- Write clean code

Requirements:
- 3+ years experience
- Python, JavaScript skills
- Bachelor's degree""",
            
            "company_profile": """Enter company information here...

Example:
Company: Tech Corp
Industry: Software Technology

Mission: Create innovative solutions
Values: Innovation, Collaboration, Excellence
Culture: Remote-first, flexible hours""",
            
            "resume": """Enter resume content here...

Example:
John Doe
Software Engineer
john@email.com

Experience:
- Software Engineer at ABC Corp (2021-Present)
- Junior Developer at XYZ Inc (2019-2021)

Education:
- BS Computer Science, State University

Skills: Python, JavaScript, React""",
            
            "response": "Type your response to the interview question here..."
        }
        
        placeholder = placeholders.get(doc_type, f"Enter {doc_type} content here...")
        
        text_content = st.text_area(
            f"Enter {label} content:",
            height=250,
            key=text_key,
            placeholder=placeholder,
            help=f"Type or paste your {doc_type} content directly"
        )
        
        if text_content and text_content.strip():
            content = text_content.strip()
            st.success("✅ Text content received!")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Characters", len(content))
            with col2:
                st.metric("Words", len(content.split()))
            with col3:
                st.metric("Lines", len(content.split('\n')))
    
    return content

# Backward compatibility
class DocumentParser:
    """Document parser class for compatibility"""
    
    @staticmethod
    def parse_pdf(uploaded_file):
        return parse_pdf(uploaded_file)
    
    @staticmethod
    def parse_docx(uploaded_file):
        return parse_docx(uploaded_file)
    
    @staticmethod
    def parse_txt(uploaded_file):
        return parse_txt(uploaded_file)
    
    @staticmethod
    def transcribe_audio(audio_bytes):
        return transcribe_audio(audio_bytes)

def display_voice_status():
    """Display voice capabilities in sidebar"""
    st.sidebar.write("### 🎤 Voice Capabilities")
    
    capabilities = check_voice_capabilities()
    
    if capabilities['voice_available']:
        st.sidebar.success(f"✅ Voice Input Available ({capabilities['voice_method']})")
        st.sidebar.write(f"🗣️ Speech Recognition: {'✅' if capabilities['speech_recognition'] else '❌'}")
    else:
        st.sidebar.error("❌ Voice Input Not Available")
        
        if st.sidebar.button("📦 Show Installation Guide"):
            st.sidebar.info("Check the Voice Input tab for installation instructions")

def main():
    """Test the enhanced parser with voice support"""
    st.title("📄 Enhanced Document Parser with Voice Support")
    
    # Display voice status
    display_voice_status()
    
    # Show capabilities
    capabilities = check_voice_capabilities()
    
    st.write("### 🔧 Available Features:")
    col1, col2 = st.columns(2)
    
    with col1:
        st.write(f"📄 PDF Parsing: {'✅' if PDF_AVAILABLE else '❌'}")
        st.write(f"📄 DOCX Parsing: {'✅' if DOCX_AVAILABLE else '❌'}")
        st.write(f"📄 TXT Parsing: ✅")
    
    with col2:
        st.write(f"🎤 Voice Recording: {'✅' if capabilities['voice_available'] else '❌'}")
        st.write(f"🗣️ Speech Recognition: {'✅' if capabilities['speech_recognition'] else '❌'}")
        if capabilities['voice_available']:
            st.write(f"📦 Method: {capabilities['voice_method']}")
    
    # Installation warnings
    if not PDF_AVAILABLE:
        st.warning("📦 Install PyPDF2 for PDF support: `pip install PyPDF2`")
    
    if not DOCX_AVAILABLE:
        st.warning("📦 Install python-docx for DOCX support: `pip install python-docx`")
    
    # Test document input
    st.write("---")
    st.write("### 🧪 Test Document Input")
    
    doc_type = st.selectbox(
        "Select document type:", 
        ["job_posting", "company_profile", "resume", "response"]
    )
    
    content = get_document_input(
        label=doc_type.replace('_', ' ').title(),
        doc_type=doc_type,
        key_suffix="test"
    )
    
    if content:
        st.write("### ✅ Content Received!")
        
        # Show content stats
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Characters", len(content))
        with col2:
            st.metric("Words", len(content.split()))
        with col3:
            st.metric("Sentences", len([s for s in content.split('.') if s.strip()]))
        
        # Show full content
        with st.expander("📄 Full Content", expanded=False):
            st.text_area("Content", content, height=300, disabled=True)
        
        # Test voice-only input
        if capabilities['voice_available']:
            st.write("---")
            st.write("### 🎤 Test Voice-Only Input")
            
            _, voice_text = voice_input_interface("voice_test", height=100)
            
            if voice_text:
                st.write("**Voice Input Result:**")
                st.text_area("Voice Text", voice_text, height=100, disabled=True)

if __name__ == "__main__":
    main()